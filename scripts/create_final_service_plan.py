from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\dev\InnerDerma\docs\AAC_WHS_AI_피부_사후관리_서비스_최종_기획서.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE, DARK_BLUE, LIGHT_BLUE = "2E74B5", "1F4D78", "EAF3F9"
PALE, GRAY, DARK, GOLD = "F4F6F9", "5F6B76", "1E2933", "C18B36"


def font(run, size=11, bold=None, color=DARK, italic=None):
    run.font.name = "Malgun Gothic"
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{key}"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tcpr.append(node)


def cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcpr.append(mar)


def geometry(table, widths, indent=120):
    table.autofit = False
    tblpr = table._tbl.tblPr
    tw = tblpr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tw.getparent() is None:
        tblpr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ti = tblpr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if ti.getparent() is None:
        tblpr.append(ti)
    ti.set(qn("w:w"), str(indent))
    ti.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tcw.getparent() is None:
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell_margins(cell)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def table(doc, headers, rows, widths, size=8.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, PALE)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), size, True, DARK_BLUE)
    repeat_header(t.rows[0])
    for values in rows:
        cells = t.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            font(p.add_run(str(value)), size)
    geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    for key, value in (("val", "single"), ("sz", "18"), ("space", "8"), ("color", BLUE)):
        left.set(qn(f"w:{key}"), value)
    border.append(left)
    ppr.append(border)
    font(p.add_run(label + "  "), 10.5, True, DARK_BLUE)
    font(p.add_run(text), 10.5)


def bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        font(p.add_run(item), 10.5)


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        font(p.add_run(item), 10.5)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(paragraph.add_run("Page "), 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance, sec.footer_distance = Inches(0.45), Inches(0.45)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Malgun Gothic", Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    s = doc.styles[name]
    s.font.name, s.font.size, s.font.bold = "Malgun Gothic", Pt(size), True
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True

hp = sec.header.paragraphs[0]
font(hp.add_run("AAC × WHS | AI 피부 사후관리 서비스 최종 기획서"), 8.5, True, GRAY)
page_number(sec.footer.paragraphs[0])

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(110)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("FINAL SERVICE PLANNING DOCUMENT"), 10, True, GOLD)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(18), Pt(10)
font(p.add_run("AAC 해외 고객\nAI 피부 사후관리 서비스"), 28, True, DARK_BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
font(p.add_run("귀가 후 저녁부터 다음 날 아침까지 이어지는 생활 주기형 케어"), 13.5, color=BLUE)
callout(doc, "핵심 정의", "WHS의 초기 피부 데이터를 바탕으로, 귀국 후 촬영한 사진과 자가 상태를 분석해 당일 저녁 집중 관리와 다음 날 아침 기본 관리를 하나의 케어 사이클로 제공하는 서비스입니다.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
font(p.add_run("최종 기획·기능안 | 2026.08.16"), 10, True, GRAY)
doc.add_page_break()

doc.add_heading("1. 서비스 개요", 1)
doc.add_paragraph("본 서비스는 AAC가 운영하는 Wellness House Seoul(WHS)에서 피부 분석을 받은 해외 고객의 귀국 후 관리를 지원하는 AI 기반 사후관리 기능이다. WHS가 제공한 초기 피부 상태와 타입을 기준 데이터로 사용하고, 사용자가 귀국 후 촬영한 피부 사진과 자가 상태를 반영하여 생활 흐름에 맞는 케어 솔루션을 생성한다.")
callout(doc, "서비스 포지셔닝", "WHS 앱을 대체하는 별도 서비스가 아니라, WHS의 초기 루틴을 해외 고객의 귀국 후 일상까지 확장하는 동적 사후관리 기능입니다.")

doc.add_heading("2. 기획 배경과 해결 과제", 1)
doc.add_paragraph("WHS는 AI 피부 분석과 초기 홈케어 루틴 및 제품 추천을 제공한다. 그러나 해외 고객은 귀국 후 전문 장비로 피부를 반복 측정하기 어렵고, 최초 루틴만으로는 외부 활동·세안·수면·시술 회복 과정에서 변하는 피부 상태를 계속 반영하기 어렵다.")
bullets(doc, [
    "귀국 후 피부 변화와 이전 상태의 차이를 확인하기 어렵다.",
    "사용자마다 귀가·수면·기상 시간이 달라 고정 시각 루틴이 실제 생활과 맞지 않는다.",
    "시술 후 통증·열감·당김 등 사진으로 보이지 않는 상태를 함께 관리해야 한다.",
    "피부 변화와 제품 사용 이력을 반영해 루틴과 추천을 지속적으로 조정할 필요가 있다.",
])

doc.add_heading("3. 사용자 유형", 1)
table(doc, ["유형", "조건", "제공 목적"], [
    ("피부 분석 사용자", "WHS에서 피부 분석만 진행", "피부 타입·상태 기반 일상 홈케어"),
    ("피부 분석·시술 사용자", "WHS 피부 분석 후 AAC 계열 클리닉 시술", "기본 홈케어에 시술 관리·주의사항 추가"),
], [2200, 3400, 3760])
callout(doc, "범위", "피부 분석 없이 시술만 받은 사용자는 이번 서비스 대상에서 제외합니다.")

doc.add_heading("4. 핵심 차별점", 1)
table(doc, ["구분", "WHS의 공개 기능", "제안 서비스"], [
    ("주요 시점", "오프라인 방문 전후", "해외 귀국 후 반복 관리"),
    ("분석 데이터", "전문 기기 기반 초기 분석", "스마트폰 사진과 체감 증상 누적"),
    ("루틴", "초기 맞춤 루틴", "촬영 시점마다 저녁·다음 날 아침 루틴 생성"),
    ("변화 추적", "공개 정보상 상세 확인 어려움", "이전 분석과 개선·유지·악화 비교"),
    ("제품 추천", "초기 분석 기반 큐레이션", "현재 상태와 주간 사용 이력에 따라 갱신"),
], [1400, 3440, 4520], 8.5)

doc.add_heading("5. 케어 사이클", 1)
doc.add_paragraph("케어 솔루션은 자정 기준 하루 단위로 묶지 않는다. 사용자의 귀가 후 세안부터 다음 날 기상 후 세안까지를 하나의 케어 사이클로 구성한다. 특정 시각을 강제하지 않고 사용자의 생활 행동을 기준으로 제공한다.")
callout(doc, "케어 사이클", "사진 촬영일의 저녁 집중 관리 + 다음 날 기상 후 아침 기본 관리")
table(doc, ["케어 사이클", "저녁 관리", "아침 관리"], [
    ("17~18일", "17일 귀가 후", "18일 기상 후"),
    ("18~19일", "18일 귀가 후", "19일 기상 후"),
    ("19~20일", "19일 귀가 후", "20일 기상 후"),
], [2800, 3280, 3280])
doc.add_paragraph("같은 날짜에는 이전 사이클의 아침 루틴과 새 사이클의 저녁 루틴이 함께 존재할 수 있다. 예를 들어 18일 아침은 17~18일 사이클에, 18일 저녁은 18~19일 사이클에 속한다.")

doc.add_heading("6. 전체 사용자 흐름", 1)
steps(doc, [
    "WHS에서 피부 상태와 피부 타입을 분석하고 초기 루틴을 제공받는다.",
    "필요한 사용자는 AAC 계열 피부 클리닉에서 시술을 받는다.",
    "WHS 앱 계정에 피부 분석 및 시술 기록을 연결한다.",
    "귀국 후 외출에서 돌아와 세안한 다음 하루 한 번 피부 사진을 촬영한다.",
    "통증·열감·당김 등 사진으로 알 수 없는 상태를 직접 입력한다.",
    "AI가 사진을 분석하고 안전·시술·제품 사용 규칙을 적용한다.",
    "LLM이 당일 저녁과 다음 날 아침의 케어 솔루션을 생성한다.",
    "다음 촬영 시 새 사이클을 만들고, 미촬영 시 기존 솔루션을 이어서 제공한다.",
])

doc.add_heading("7. 피부 사진 촬영", 1)
doc.add_heading("7.1 촬영 정책", 2)
bullets(doc, [
    "사용자의 현지 날짜를 기준으로 하루 한 번만 유효 촬영을 허용한다.",
    "외출 후 귀가하여 세안을 마치고, 화장품을 바르기 전에 촬영한다.",
    "가능한 한 같은 조명·거리·각도를 유지하고 필터와 피부 보정을 사용하지 않는다.",
    "화질·조명·얼굴 위치 검사를 통과한 사진만 촬영 횟수에 포함한다.",
])
doc.add_heading("7.2 촬영 날짜 데이터", 2)
table(doc, ["필드", "의미"], [
    ("capturedAt", "실제 촬영 일시"),
    ("captureDate", "사용자의 현지 촬영 날짜"),
    ("timezone", "촬영 당시 사용자 시간대"),
    ("imageUrl", "저장된 피부 사진 위치"),
    ("analysisId", "사진으로 생성된 AI 분석 결과"),
], [2600, 6760])
doc.add_paragraph("동일 사용자와 동일 현지 날짜의 유효 촬영은 하나만 허용한다. 서버 날짜가 아닌 사용자의 현지 날짜를 기준으로 제한하며, 품질 검사에 실패한 사진은 제한에 포함하지 않는다.")

doc.add_heading("8. AI 사진 분석과 자가 상태", 1)
doc.add_heading("8.1 AI 사진 분석", 2)
table(doc, ["분석 항목", "활용"], [
    ("색소 불균형", "현재 상태와 이전 촬영 대비 변화"),
    ("모공", "현재 상태와 이전 촬영 대비 변화"),
    ("주름", "현재 상태와 이전 촬영 대비 변화"),
    ("홍조", "현재 상태와 이전 촬영 대비 변화"),
    ("피부결", "현재 상태와 이전 촬영 대비 변화"),
], [2600, 6760])
doc.add_paragraph("결과는 IMPROVED, STABLE, WORSENED, NEEDS_ATTENTION으로 구분한다. 이는 피부 상태 변화를 확인하는 보조 정보이며 의료 진단을 의미하지 않는다.")
doc.add_heading("8.2 자가 상태", 2)
bullets(doc, ["통증·열감·당김·건조함", "가려움·붓기·피부 벗겨짐", "트러블·진물·출혈", "피부 장벽 손상 체감"])
doc.add_paragraph("각 항목은 없음·약함·보통·심함으로 입력한다. 사진을 촬영하지 않은 날에도 위험 증상을 입력하면 기존 솔루션 승계보다 안전 안내를 우선한다.")

doc.add_heading("9. 케어 솔루션 생성", 1)
doc.add_heading("9.1 입력 데이터", 2)
table(doc, ["우선순위", "데이터", "역할"], [
    ("기본", "WHS 피부 타입·초기 피부 상태", "장기 기준점"),
    ("현재", "최신 사진 분석·자가 상태·이전 변화", "현재 관리 목적 결정"),
    ("사용", "보유 제품·주간 사용 횟수·마지막 사용일", "제품 배치와 제한"),
    ("추가", "시술 종류·시술일·경과일·주의사항", "관리·금지 조건 추가"),
], [1200, 4400, 3760])
doc.add_heading("9.2 생성 방식", 2)
doc.add_paragraph("케어 규칙과 사용자 데이터를 시스템 프롬프트 및 구조화된 컨텍스트로 LLM에 전달한다. 이는 별도의 모델 학습이나 파인튜닝이 아니라 규칙 기반 컨텍스트 제공 방식이다.")
steps(doc, [
    "코드에서 위험 증상과 필수 안전 규칙을 확인한다.",
    "사용 가능한 제품 후보와 금지 성분을 결정한다.",
    "시술 경과일 및 주간 사용 제한을 적용한다.",
    "LLM이 허용된 범위 안에서 아침·저녁 루틴과 사용자 설명을 생성한다.",
    "출력 형식과 금지 조건을 다시 검증한 후 사용자에게 제공한다.",
])

doc.add_heading("10. 저녁·아침 케어", 1)
table(doc, ["구분", "저녁 집중 관리", "다음 날 아침 기본 관리"], [
    ("기준", "귀가 후 세안·사진·자가 상태", "전날 저녁에 확인된 최신 상태"),
    ("목적", "진정·보습·장벽·회복", "수분 유지·보호·자외선 차단"),
    ("깊이", "재생 앰플 등 집중 관리 가능", "간단하고 보호 중심의 스킨케어"),
    ("안내", "제품 순서·용량·금지 성분·취침 전 수칙", "세안·가벼운 보습·외출 전 수칙"),
], [1400, 3980, 3980], 8.5)

doc.add_heading("11. 촬영·미촬영 처리", 1)
doc.add_heading("11.1 촬영한 경우", 2)
callout(doc, "예시", "17일 촬영 → 17일 저녁 케어 + 18일 아침 케어 → 17~18일 신규 케어 사이클")
doc.add_heading("11.2 촬영하지 않은 경우", 2)
doc.add_paragraph("새로운 사진이 없으면 가장 최근에 생성된 솔루션을 현재 사이클에 그대로 적용한다. 새로운 분석을 한 것처럼 표시하지 않고 원본 촬영일과 승계 상태를 함께 보여준다.")
table(doc, ["케어 사이클", "촬영 여부", "기준 촬영일", "제공 방식"], [
    ("17~18일", "촬영", "17일", "신규 분석·솔루션"),
    ("18~19일", "미촬영", "17일", "이전 솔루션 유지"),
    ("19~20일", "촬영", "19일", "신규 분석·솔루션"),
], [2200, 1800, 2200, 3160])
callout(doc, "사용자 안내", "오늘은 새로운 촬영 기록이 없어 17일 피부 상태를 기준으로 오늘 저녁과 다음 날 아침 루틴을 제공하고 있어요.", "FFF7E8")

doc.add_heading("12. 캘린더와 케어 기록", 1)
doc.add_paragraph("캘린더는 날짜별로 아침·저녁 루틴, 적용 사이클, 기준 촬영일과 생성 유형을 표시한다. 한 날짜의 아침과 저녁이 서로 다른 사이클에 속할 수 있으므로 시간대별 기록을 분리한다.")
table(doc, ["날짜", "시점", "적용 사이클", "기준 촬영일", "상태"], [
    ("17일", "저녁", "17~18일", "17일", "NEW_ANALYSIS"),
    ("18일", "아침", "17~18일", "17일", "신규 분석 기반"),
    ("18일", "저녁", "18~19일", "17일", "CARRIED_FORWARD"),
    ("19일", "아침", "18~19일", "17일", "기존 솔루션 유지"),
], [1200, 1200, 2100, 2100, 2760], 8.2)
bullets(doc, [
    "NEW_ANALYSIS: 당일 유효 사진으로 새 분석과 솔루션 생성",
    "CARRIED_FORWARD: 최근 솔루션을 현재 케어 사이클에 이어서 제공",
    "위험 증상 발생일은 별도 경고 상태로 표시",
])

doc.add_heading("13. 주간 제품 예외와 제품 추천", 1)
doc.add_paragraph("케어 솔루션은 매 사이클 제공하지만 모든 제품과 영양제를 매일 사용하지 않는다. 제품별 주간 최대 횟수·마지막 사용일·연속 사용 제한·시술 후 제한을 코드에서 관리한다.")
table(doc, ["관리 항목", "빈도 예시", "적용"], [
    ("보습제", "매일", "아침·저녁"),
    ("자외선 차단제", "매일", "아침·외출 전"),
    ("진정 마스크", "주 2회", "저녁"),
    ("각질 관리", "주 1회", "상태·시술 조건 확인 후 저녁"),
    ("기능성 앰플", "주 2~3회", "저녁 집중 관리"),
    ("영양제", "제품 기준", "지정 요일·횟수"),
], [2800, 2100, 4460])
doc.add_paragraph("제품은 현재 피부 상태와 관리 목적을 먼저 설명한 뒤, 관련 성분이 포함된 Pith 및 WHS Store 제품을 추천한다. 구매한 제품은 사용 가능 여부를 확인해 아침·저녁·주간 루틴에 자동 배치한다.")

doc.add_heading("14. 시술 사용자 처리", 1)
doc.add_paragraph("시술 정보는 피부 상태를 대신 판단하는 데이터가 아니라 기본 케어 솔루션에 관리 방법과 주의사항을 추가하는 조건으로 사용한다.")
bullets(doc, [
    "시술 종류·시술일·시술 후 경과일을 확인한다.",
    "시술 후 피해야 할 성분·제품·행동을 적용한다.",
    "사진 분석으로 눈에 보이는 변화를 추적하되 정상 회복이나 부작용을 확정하지 않는다.",
    "통증·열감·붓기 등 위험 신호가 있으면 시설 또는 의료기관 문의를 안내한다.",
])

doc.add_heading("15. 핵심 데이터 모델", 1)
table(doc, ["데이터", "핵심 필드"], [
    ("SkinCapture", "userId, capturedAt, captureDate, timezone, imageUrl, analysisId"),
    ("SkinAnalysis", "captureId, 항목별 점수, resultStatus, analyzedAt"),
    ("SelfCheck", "사용자 증상별 강도, checkedAt"),
    ("CareCycle", "cycleStartDate, cycleEndDate, eveningRoutineDate, morningRoutineDate"),
    ("솔루션 출처", "sourceCaptureId, sourceAnalysisId, sourceSolutionId, generationType"),
    ("ProductUsage", "productId, usedAt, weeklyCount, lastUsedAt"),
], [2600, 6760], 8.5)

doc.add_heading("16. 필수 기능과 선택 기능", 1)
doc.add_heading("16.1 MVP 필수 기능", 2)
bullets(doc, [
    "WHS 계정 연결을 표현하는 로그인과 사용자 코드",
    "피부 분석 사용자와 피부 분석·시술 사용자 구분",
    "WHS 초기 피부 및 시술 더미 데이터 연결",
    "현지 날짜 기준 하루 한 번 사진 촬영과 품질 검사",
    "AI 사진 분석과 자가 상태 확인",
    "저녁·다음 날 아침 케어 사이클 생성",
    "미촬영 시 이전 솔루션 승계 및 원본 기준일 표시",
    "주간 제품 사용 제한과 시술 주의사항 적용",
    "제품 추천·상세·구매 연결",
    "캘린더 및 케어 기록",
])
doc.add_heading("16.2 선택 기능", 2)
bullets(doc, ["피부 변화 그래프", "촬영·루틴·재구매 알림", "다국어 UI", "시설 문의 화면", "마이페이지·콘텐츠", "직원용 등록 화면"])

doc.add_heading("17. 안전 및 표현 원칙", 1)
table(doc, ["지양 표현", "권장 표현"], [
    ("AI 피부 진단", "AI 피부 상태 분석"),
    ("치료·처방", "케어 솔루션·추천·관리 안내"),
    ("성분 적합도", "현재 관리 목적에 대한 제품 추천도"),
    ("정상입니다", "일반적인 경과일 수 있습니다"),
], [3600, 5760])
bullets(doc, [
    "AI 결과는 전문적인 의료 판단을 대체하지 않는다.",
    "제품은 개인에 따라 자극이나 알레르기를 유발할 수 있다.",
    "위험 증상은 일반 케어보다 안전 안내를 우선한다.",
    "영양제는 복용 약물·알레르기·기저질환을 고려한다.",
])

doc.add_heading("18. 비즈니스 모델", 1)
doc.add_paragraph("서비스는 WHS 앱의 무료 사후관리 기능으로 제공한다. AAC는 사용자의 최신 피부 상태와 관리 목적에 맞는 Pith 및 WHS Store 제품을 추천하고, 구매 제품을 루틴에 연결해 첫 구매와 재구매를 만든다.")
callout(doc, "비즈니스 흐름", "무료 귀국 후 관리 → WHS 앱 재방문 → 피부 상태 기반 제품 추천 → 구매·루틴 연결 → 사용 기록 → 재구매·AAC 재방문")

doc.add_heading("19. 최종 서비스 정의", 1)
callout(doc, "최종 정의", "WHS에서 피부 분석을 받은 해외 고객이 귀국 후 외출에서 돌아와 세안한 피부를 하루 한 번 촬영하고 자가 상태를 입력하면, AI가 당일 저녁 집중 관리와 다음 날 아침 기본 관리를 하나의 케어 사이클로 생성합니다. 다음 날 새 사진이 있으면 새 사이클을 만들고, 촬영하지 않으면 최근 솔루션을 원본 촬영일과 함께 이어서 제공합니다. 시술 사용자는 기본 솔루션에 시술 관리와 주의사항을 추가하며, 매일 사용할 수 없는 제품과 영양제는 주간 사용 이력을 기준으로 예외 처리합니다.")

doc.add_heading("20. 공식 정보 및 검증 과제", 1)
doc.add_heading("20.1 참고 자료", 2)
bullets(doc, [
    "Wellness House Seoul: https://www.wellnesshouse-seoul.com/ko",
    "WHS Store 및 맞춤 루틴: https://www.wellnesshouse-seoul.com/ko/wellness",
    "WHS 공식 앱: https://apps.apple.com/kr/app/id6760689802",
    "AAC 클리닉: https://anti-agingclub.kr/clinic",
    "AAC Pith: https://anti-agingclub.kr/55",
])
doc.add_heading("20.2 추가 검증 과제", 2)
bullets(doc, [
    "WHS 앱의 실제 사후관리 기능과 화면 범위",
    "AAC·WHS·클리닉의 실제 데이터 구조와 연동 가능성",
    "사용할 오픈소스 피부 분석 모델의 정확도와 라이선스",
    "시술별 관리·금지·위험 규칙에 대한 전문가 검수",
    "사용자 시간대 변경 및 국제 날짜 변경 시 촬영 정책",
])

doc.core_properties.title = "AAC WHS AI 피부 사후관리 서비스 최종 기획서"
doc.core_properties.subject = "귀국 후 생활 주기형 AI 피부 케어 사이클"
doc.core_properties.author = "InnerDerma Team"
doc.core_properties.keywords = "AAC, WHS, 피부 분석, 케어 사이클, 사후관리, Pith"
doc.save(OUT)
print(OUT)
