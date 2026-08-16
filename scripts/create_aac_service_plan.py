from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\dev\InnerDerma\docs\AAC_WHS_AI_피부_사후관리_서비스_기획서.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF3F9"
PALE = "F4F6F9"
GRAY = "5F6B76"
DARK = "1E2933"
WHITE = "FFFFFF"
GOLD = "C18B36"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=11, bold=None, color=DARK, italic=None):
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(label + "  ")
    set_run_font(r, 10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, 10.5, color=DARK)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, font_size, bold=True, color=DARK_BLUE)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run_font(r, font_size, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.45)
section.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    st = styles[name]
    st.font.name = "Malgun Gothic"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run("AAC × WHS | AI 피부 사후관리 서비스")
set_run_font(r, 8.5, bold=True, color=GRAY)
footer = section.footer.paragraphs[0]
page_number(footer)

# Cover: editorial_cover pattern
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(112)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SERVICE PLANNING DOCUMENT")
set_run_font(r, 10, bold=True, color=GOLD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("AAC 해외 고객\nAI 피부 사후관리 서비스")
set_run_font(r, 28, bold=True, color=DARK_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run("WHS의 첫 루틴을, 귀국 후 피부 변화에 맞춰 계속 진화시키다")
set_run_font(r, 13.5, color=BLUE)

add_callout(doc, "핵심 정의", "WHS의 초기 피부 분석과 관리 루틴을 기반으로, 귀국 후 사용자의 사진·자가 상태·시술 경과를 반영해 케어 솔루션을 지속적으로 갱신하는 확장형 사후관리 기능입니다.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(52)
r = p.add_run("해커톤 기획안 | 2026.08.15")
set_run_font(r, 10, bold=True, color=GRAY)

doc.add_page_break()

doc.add_heading("1. 서비스 개요", level=1)
doc.add_paragraph("본 서비스는 AAC가 운영하는 Wellness House Seoul(WHS)의 오프라인 피부 분석 및 초기 맞춤 루틴을 해외 고객의 귀국 후 일상까지 확장하는 AI 사후관리 기능이다. 사용자가 스마트폰으로 촬영한 피부 사진과 직접 입력한 체감 증상을 WHS의 초기 데이터 및 시술 기록과 함께 분석하여, 피부 변화에 맞는 케어 루틴·제품 사용 방식·관리 주기를 지속적으로 조정한다.")
add_callout(doc, "서비스 한 줄 정의", "WHS가 고객의 첫 번째 웰니스 루틴을 설계한다면, 우리는 그 루틴이 귀국 후에도 피부 변화에 맞춰 계속 진화하도록 만듭니다.")

doc.add_heading("2. 기업·서비스 맥락", level=1)
doc.add_paragraph("AAC는 클리닉, 피부·건강 데이터, 홈케어 제품과 오프라인 웰니스 경험을 연결하는 안티에이징·웰니스 기업이다. WHS는 AI 기반 분석, 맞춤형 루틴, 제품 큐레이션, 클리닉, 웰니스 프로그램과 식음료 경험을 연결하는 통합 공간 및 앱이다.")
add_table(doc, ["구분", "확인된 역할"], [
    ("AAC", "클리닉·웰니스 공간·제품·디지털 경험을 연결하는 기업"),
    ("WHS", "AI 피부·컨디션 분석, 초기 루틴, 제품·프로그램 추천 및 앱 연동"),
    ("DERNA / AMRED", "AAC의 피부 관련 클리닉 브랜드"),
    ("Pith", "AMRED의 노하우를 기반으로 한 AAC 스킨케어·홈케어 브랜드"),
    ("WHS Store", "Pith를 포함한 스킨케어·뷰티 디바이스·헬스케어 제품 큐레이션"),
], [1800, 7560])

doc.add_heading("3. 문제 정의", level=1)
doc.add_paragraph("WHS는 공개 자료상 이미 AI 분석 결과를 리포트, 초기 홈케어 루틴 및 제품 추천으로 연결한다. 따라서 'WHS는 분석만 하고 우리는 솔루션을 제공한다'는 설명은 차별점으로 사용할 수 없다. 해결해야 할 실제 문제는 해외 고객이 귀국한 뒤 WHS 전문 장비로 피부를 다시 측정하기 어렵고, 최초 루틴이 이후의 피부 변화와 체감 증상을 따라가지 못할 수 있다는 관리 공백이다.")
add_bullets(doc, [
    "귀국 후 현재 피부 상태와 이전 상태의 차이를 확인하기 어렵다.",
    "시술 후 회복 과정에서 통증·열감·당김 같은 주관적 증상을 함께 관리하기 어렵다.",
    "피부가 개선되거나 악화되어도 최초 제품 사용법과 관리 주기가 그대로 남을 수 있다.",
    "오프라인 방문이 끝나면 AAC와 해외 고객의 장기적인 접점이 약해질 수 있다.",
])

doc.add_heading("4. 기존 WHS와의 차별점", level=1)
add_table(doc, ["구분", "WHS의 공개 기능", "제안 기능"], [
    ("주요 시점", "오프라인 방문 전후", "귀국 후 반복 관리"),
    ("피부 데이터", "전문 기기 기반 초기 분석", "스마트폰 사진과 체감 증상 누적"),
    ("루틴", "초기 맞춤 루틴", "변화에 따라 내용·용량·기간 재조정"),
    ("상태 비교", "공개 정보상 상세 확인 어려움", "이전 기록과 개선·유지·악화 비교"),
    ("시술 관리", "사후관리 제공", "경과일·사진·자가 상태를 함께 반영"),
    ("제품 추천", "초기 분석 기반 큐레이션", "최신 상태와 회복 단계에 따라 갱신"),
], [1500, 3450, 4410], 8.7)
add_callout(doc, "포지셔닝", "별도 경쟁 앱이 아니라 WHS 앱에 추가되는 '해외 고객 귀국 후 동적 사후관리 모듈'로 제안합니다.", fill="FFF7E8")

doc.add_heading("5. 사용자 유형", level=1)
doc.add_heading("5.1 피부 분석만 받은 사용자", level=2)
doc.add_paragraph("WHS에서 피부 상태와 피부 타입을 분석받았지만 피부 관련 시술은 받지 않은 사용자다. 피부 고민과 귀국 후 변화에 맞는 일상 홈케어를 제공한다.")
add_bullets(doc, ["일상적인 피부 상태 관리", "아침·저녁 홈케어 루틴", "피부 변화 추적", "현재 관리 목적에 맞는 제품 추천"])
doc.add_heading("5.2 피부 분석과 시술을 모두 받은 사용자", level=2)
doc.add_paragraph("WHS에서 피부 분석을 받은 후 DERNA 또는 AAC 계열 피부 클리닉에서 시술까지 받은 사용자다. 초기 피부 상태와 시술 종류·경과일을 함께 반영한 회복 관리가 필요하다.")
add_bullets(doc, ["시술 후 회복 과정 확인", "시술 후 사용 가능한 제품 안내", "피해야 할 성분과 행동 안내", "위험 신호 발생 시 시설 문의 안내"])
add_callout(doc, "범위 제한", "피부 분석 없이 시술만 받은 사용자는 이번 서비스 범위에서 제외합니다.")

doc.add_heading("6. 전체 사용자 여정", level=1)
doc.add_heading("6.1 피부 분석 사용자", level=2)
add_steps(doc, [
    "WHS에서 AI 피부 분석을 받고 초기 리포트와 루틴을 확인한다.",
    "WHS 앱 계정에 분석 결과를 연결하고 추천 제품을 확인·구매한다.",
    "귀국 후 스마트폰으로 피부 사진을 촬영하고 자가 상태를 입력한다.",
    "초기 데이터와 최신 상태를 비교해 새로운 아침·저녁 루틴을 제공받는다.",
    "루틴을 수행하고 지정된 날짜에 다시 촬영한다.",
    "개선·유지·악화 결과에 따라 솔루션을 유지하거나 재생성한다.",
])
doc.add_heading("6.2 피부 분석·시술 사용자", level=2)
add_steps(doc, [
    "WHS에서 피부 분석 후 AAC 계열 피부 클리닉에서 시술을 받는다.",
    "시술 당일 관리 방법과 사용 가능한 제품을 안내받는다.",
    "피부 분석·시술 기록을 WHS 앱 계정에 연결한다.",
    "귀국 후 피부 사진과 통증·열감·당김 등의 체감 증상을 입력한다.",
    "시술 경과일과 최신 상태를 종합한 회복 루틴을 제공받는다.",
    "피부 변화를 반복 확인하여 관리 강도와 제품 추천을 조정한다.",
])

doc.add_heading("7. 데이터 구성", level=1)
add_table(doc, ["데이터군", "주요 항목", "구현 방식"], [
    ("WHS 초기 분석", "피부 타입, 수분·유분, 탄력, 주름, 색소, 홍조, 모공, 피부결", "더미 데이터"),
    ("시술 기록", "시설, 종류, 날짜, 부위, 경과일, 주의 사항", "더미 데이터"),
    ("귀국 후 사진", "색소 불균형, 모공, 주름, 홍조, 피부결", "사용자 입력·오픈소스 모델"),
    ("자가 상태", "통증, 열감, 당김, 건조함, 가려움, 붓기 등", "사용자 입력"),
    ("루틴 기록", "제품 사용, 행동 수행, 촬영일, 다음 확인일", "서비스 저장"),
    ("제품 정보", "브랜드, 성분, 기능, 용량, 사용법, 가격", "공개 정보 또는 더미"),
], [1700, 4860, 2800], 8.5)
doc.add_paragraph("해커톤에서는 AAC·WHS의 실제 API와 데이터 스키마를 사용할 수 없으므로, 시설 코드 또는 사용자 코드를 입력하면 미리 설계한 피부 분석·시술 기록이 계정에 연결되는 방식으로 시연한다.")

doc.add_heading("8. 핵심 기능", level=1)
doc.add_heading("8.1 WHS 계정 및 기록 연결", level=2)
add_bullets(doc, ["이름·연락처·국가·언어 확인", "최근 방문 시설과 연결 코드 입력", "사용자 유형 자동 구분", "초기 피부 분석 및 시술 더미 기록 호출"])

doc.add_heading("8.2 AI 사진 분석", level=2)
doc.add_paragraph("사용자가 스마트폰으로 촬영한 얼굴 사진을 오픈소스 모델로 분석한다. 결과는 현재 상태를 단독으로 보여주는 데 그치지 않고 이전 촬영 결과와 비교한다.")
add_table(doc, ["분석 항목", "제공 결과"], [
    ("색소 불균형", "현재 상태·이전 대비 변화"), ("모공", "현재 상태·이전 대비 변화"),
    ("주름", "현재 상태·이전 대비 변화"), ("홍조", "현재 상태·이전 대비 변화"),
    ("피부결", "현재 상태·이전 대비 변화"),
], [2500, 6860])

doc.add_heading("8.3 자가 상태 확인", level=2)
doc.add_paragraph("사진으로 확인하기 어려운 주관적인 피부 상태를 사용자가 직접 입력한다. 기본 입력 단계는 '없음·약함·보통·심함'으로 구성한다.")
add_bullets(doc, ["통증·열감·당김·건조함", "가려움·붓기·피부 벗겨짐", "트러블·진물·출혈", "피부 장벽 손상 체감"])

doc.add_heading("8.4 피부 상태 종합 분석", level=2)
doc.add_paragraph("초기 피부 상태, 시술 여부와 경과일, 최신 사진 분석, 자가 상태, 이전 기록 및 루틴 수행 여부를 종합하여 현재의 관리 목적을 결정한다.")
add_bullets(doc, ["피부 진정", "수분 유지", "피부 장벽 관리", "홍조·색소·피부결 관리", "자외선 차단", "피부 자극 최소화"])

doc.add_heading("9. AI 케어 솔루션", level=1)
doc.add_heading("9.1 오늘의 피부 상태", level=2)
add_callout(doc, "일상 관리 예시", "현재 피부는 이전보다 수분 상태가 낮고 당김이 확인되었습니다. 오늘은 수분 유지와 피부 장벽 관리에 집중해주세요.")
add_callout(doc, "시술 관리 예시", "시술 후 3일이 경과했습니다. 홍조는 이전보다 감소했지만 당김이 남아 있으므로 진정과 보습 루틴을 유지해주세요.")

doc.add_heading("9.2 아침·저녁 루틴", level=2)
add_table(doc, ["구분", "솔루션 구성"], [
    ("아침", "세안 방식, 제품 순서, 권장 사용량, 자외선 차단, 행동 수칙"),
    ("저녁", "클렌징, 진정·보습, 제품 순서, 피해야 할 성분·행동"),
    ("음식·생활", "수분 섭취, 자극적 음식·음주·격한 운동·사우나 제한, 수면"),
    ("영양·헬스케어", "WHS Store 제품의 추천 이유, 섭취 방법·주기·주의 사항"),
], [1800, 7560])

doc.add_heading("9.3 제품 추천", level=2)
doc.add_paragraph("제품 자체를 먼저 노출하지 않고 피부 상태, 관리 목적, 도움이 될 수 있는 성분군, 해당 성분을 포함한 WHS Store 제품의 순서로 추천 근거를 설명한다. 화장품은 Pith를 포함한 WHS Store 스킨케어 제품을, 영양제는 별도의 헬스케어 제품군을 대상으로 한다.")
add_steps(doc, ["현재 피부 상태 설명", "현재 필요한 관리 목적 도출", "관리에 활용되는 성분군 안내", "조건에 맞는 제품 추천", "사용 순서·용량·기간과 구매 버튼 제공"])

doc.add_heading("9.4 제품 추천도", level=2)
doc.add_paragraph("0~100 점수를 사용하는 경우 '성분 적합도'가 아니라 '제품 추천도'로 정의한다. 이는 현재 피부 상태와 관리 목적에 해당 제품이 얼마나 부합하는지 나타내며, 알레르기 안전성이나 치료 효과를 의미하지 않는다.")
add_table(doc, ["제품 예시", "추천도", "근거"], [
    ("진정 크림", "92", "진정·장벽 관리 목적과 일치"),
    ("수분 세럼", "86", "건조함·당김 관리에 활용 가능"),
    ("고기능성 앰플", "35", "시술 직후 자극 가능성으로 사용 보류"),
], [2700, 1400, 5260])

doc.add_heading("10. 가변적인 솔루션 주기", level=1)
doc.add_paragraph("각 관리 항목에는 행동의 빈도인 '반복 주기'와 행동을 유지하는 총 기간인 '수행 기간'을 별도로 저장한다. 전체 솔루션 유효기간은 각 항목의 수행 기간 중 가장 긴 값으로 설정한다.")
add_table(doc, ["관리 항목", "반복 주기", "수행 기간"], [
    ("진정 크림", "매일 아침·저녁", "5일"),
    ("진정 마스크", "2일마다", "6일"),
    ("자외선 차단", "매일 아침", "7일"),
    ("음주 제한", "지속", "3일"),
], [3300, 3300, 2760])
add_callout(doc, "계산 예시", "위 솔루션의 최대 수행 기간은 7일이므로 전체 유효기간과 다음 피부 확인일을 7일 후로 설정합니다.")

doc.add_heading("11. 재확인 결과와 분기", level=1)
add_table(doc, ["분기", "판단 기준", "서비스 반응"], [
    ("개선", "사진 결과 개선, 자가 상태 완화, 위험 신호 없음", "기존 루틴 유지, 개선 항목과 긍정적 메시지 제공"),
    ("유지", "이전과 유사, 자가 상태 악화 없음", "유효기간 내 기존 루틴 유지, 안정적 경과 안내"),
    ("악화", "사진 결과 또는 체감 증상 악화", "자극 제품 중단, 특정 루틴 조정, 시설 문의 안내"),
    ("위험", "심한 통증·열감·붓기, 진물·출혈·물집", "일반 솔루션보다 안전 안내 우선"),
    ("기간 만료", "전체 솔루션 유효기간 종료", "최신 데이터로 전체 솔루션 재생성"),
], [1200, 3720, 4440], 8.3)
add_callout(doc, "메시지 원칙", "개선이 확인되지 않았는데 '점점 좋아지고 있어요'라고 표현하지 않습니다. 변화가 없다면 '안정적으로 유지되고 있어요'라고 안내합니다.", fill="FFF7E8")

doc.add_heading("12. 케어 기록", level=1)
add_bullets(doc, [
    "피부 사진과 촬영 날짜", "AI 분석 결과와 이전 대비 변화", "자가 상태 입력 결과",
    "시술 후 경과일", "제공된 케어 솔루션", "제품 구매·사용 기록",
    "루틴 수행 여부", "개선·유지·악화 결과", "다음 피부 확인일",
])

doc.add_heading("13. 제품 구매와 AAC 비즈니스 연결", level=1)
doc.add_paragraph("서비스는 사용자에게 무료로 제공한다. AAC는 변화하는 피부 상태에 맞춰 WHS Store 제품을 추천하고, 구매한 제품을 케어 루틴에 자동으로 연결함으로써 첫 구매와 재구매를 유도한다.")
add_steps(doc, [
    "무료 귀국 후 사후관리로 WHS 앱 재방문을 만든다.",
    "최신 피부 상태를 근거로 제품을 추천한다.",
    "구매 제품을 아침·저녁 루틴에 자동 배치한다.",
    "사용 기록과 피부 변화를 확인한다.",
    "필요 시 추천을 변경하거나 재구매를 안내한다.",
    "장기적으로 AAC 클리닉과 WHS 재방문으로 연결한다.",
])

doc.add_heading("14. MVP 기능 범위", level=1)
doc.add_heading("14.1 필수 기능", level=2)
add_bullets(doc, [
    "WHS 계정 연결을 표현하는 로그인과 사용자 코드", "두 사용자 유형 구분",
    "초기 피부 분석 및 시술 더미 기록", "AI 사진 분석", "자가 상태 확인",
    "이전·현재 결과 비교", "개선·유지·악화·위험 분기", "아침·저녁 루틴",
    "음식·영양·행동 수칙", "반복 주기와 수행 기간", "WHS Store 제품 추천",
    "제품 상세·구매 연결", "케어 기록", "솔루션 유지·재생성",
])
doc.add_heading("14.2 선택 기능", level=2)
add_bullets(doc, ["피부 변화 그래프", "루틴 수행·재촬영·재구매 알림", "다국어 UI", "시설 문의 화면", "마이페이지와 콘텐츠 추천", "직원용 사용자 등록 화면"])

doc.add_heading("15. 안전 및 표현 원칙", level=1)
add_table(doc, ["지양 표현", "권장 표현"], [
    ("AI 피부 진단", "AI 피부 상태 분석"),
    ("치료 솔루션", "케어 솔루션"),
    ("처방", "추천·관리 안내"),
    ("성분 적합도", "제품 추천도"),
    ("이 제품은 안전합니다", "현재 관리 목적에 부합합니다"),
    ("정상입니다", "일반적인 경과일 수 있습니다"),
], [3500, 5860])
add_bullets(doc, [
    "AI 결과는 전문적인 의료 판단을 대체하지 않는다.",
    "제품은 개인에 따라 자극이나 알레르기를 유발할 수 있다.",
    "위험 증상이 있으면 시술 시설 또는 의료기관 문의를 안내한다.",
    "영양제는 복용 약물·알레르기·기저질환을 고려해야 한다.",
])

doc.add_heading("16. 핵심 성과 지표", level=1)
add_table(doc, ["목표", "지표 예시"], [
    ("귀국 후 사용 지속", "가입 후 7일·30일 재방문율, 재촬영 완료율"),
    ("케어 실행", "루틴 체크율, 솔루션 완료율"),
    ("제품 전환", "추천 제품 클릭률, 장바구니율, 구매 전환율"),
    ("반복 구매", "30일·60일 재구매율"),
    ("AAC 관계 유지", "시설 문의율, 재예약·재방문 의향"),
], [3000, 6360])

doc.add_heading("17. 발표 메시지", level=1)
add_callout(doc, "문제", "WHS는 고객의 피부를 정밀하게 분석하고 초기 루틴을 설계하지만, 해외 고객은 귀국 후 변화하는 피부를 같은 전문 장비로 다시 확인하기 어렵습니다.")
add_callout(doc, "해결", "귀국 후 사진과 통증·열감·당김 등의 체감 증상을 수집하고, WHS 초기 데이터와 비교하여 루틴과 제품 사용 방법을 지속적으로 갱신합니다.")
add_callout(doc, "차별성", "WHS가 루틴을 시작한다면, 우리는 그 루틴이 귀국 후에도 피부 변화에 맞춰 계속 진화하도록 만듭니다.")
add_callout(doc, "사업 가치", "사용자에게는 무료 사후관리를, AAC에는 제품 구매·재구매와 장기 고객 접점을 제공합니다.")

doc.add_heading("18. 검증이 필요한 사항", level=1)
add_bullets(doc, [
    "WHS 앱의 실제 화면과 세부 사후관리 범위", "WHS·클리닉의 실제 데이터 스키마와 연동 가능 여부",
    "해외 고객이 시술받는 정확한 시설 및 동선", "Pith 및 WHS Store 제품의 최신 상품·성분 데이터",
    "AI 분석에 사용할 오픈소스 모델의 상업적 이용 라이선스와 정확도", "시술별 위험 신호·관리 기준에 대한 전문가 검수",
])

doc.add_heading("19. 참고 자료", level=1)
sources = [
    "Wellness House Seoul 공식 사이트: https://www.wellnesshouse-seoul.com/ko",
    "WHS Store 및 맞춤 루틴 소개: https://www.wellnesshouse-seoul.com/ko/wellness",
    "WHS 공식 App Store 페이지: https://apps.apple.com/kr/app/id6760689802",
    "AAC 공식 클리닉 소개: https://anti-agingclub.kr/clinic",
    "AAC 공식 Pith 소개: https://anti-agingclub.kr/55",
]
add_bullets(doc, sources)

doc.core_properties.title = "AAC 해외 고객 AI 피부 사후관리 서비스 기획서"
doc.core_properties.subject = "WHS 귀국 후 동적 피부 사후관리 기능"
doc.core_properties.author = "InnerDerma Team"
doc.core_properties.keywords = "AAC, WHS, 피부 분석, 사후관리, 케어 솔루션, Pith"

doc.save(OUT)
print(str(OUT))
